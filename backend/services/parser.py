import csv
import io
import re
from typing import Optional, Tuple


def parse_uploaded_file(filename: str, content: bytes) -> list[tuple[str, str]]:
    """Returns list of (source_word, target_word) pairs."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(content)
    elif ext == "docx":
        return _parse_docx(content)
    else:
        return _parse_csv_or_txt(content)


def _strip_prefix(line: str) -> str:
    """Remove leading numbers, bullets, dashes from a line."""
    line = line.strip()
    line = re.sub(r"^\d+[.)]\s*", "", line)   # "1. " or "1) "
    line = re.sub(r"^[-•*·]\s*", "", line)     # "- " or "• "
    return line.strip()


def _split_pair(line: str) -> Optional[Tuple[str, str]]:
    """Try to split a line into (source, target) using common delimiters."""
    line = _strip_prefix(line)
    if not line:
        return None
    for sep in ["\t", " - ", " – ", " — ", ",", ";", " = ", ":"]:
        if sep in line:
            parts = line.split(sep, 1)
            src, tgt = parts[0].strip(), parts[1].strip()
            # Filter out header-like rows
            if src and tgt and src.lower() not in ("word", "woord", "dutch", "english", "translation", "term"):
                return src, tgt
    return None


def _parse_csv_or_txt(content: bytes) -> list[tuple[str, str]]:
    text = content.decode("utf-8", errors="replace")
    pairs: list[tuple[str, str]] = []

    # Try CSV reader with comma first
    try:
        reader = csv.reader(io.StringIO(text))
        candidates: list[tuple[str, str]] = []
        for row in reader:
            if len(row) >= 2:
                src = _strip_prefix(row[0].strip())
                tgt = row[1].strip()
                if src and tgt and src.lower() not in ("word", "woord", "dutch", "english"):
                    candidates.append((src, tgt))
        if candidates:
            return candidates
    except Exception:
        pass

    # Fallback: line by line with delimiter detection
    for line in text.splitlines():
        pair = _split_pair(line)
        if pair:
            pairs.append(pair)

    return pairs


def _parse_pdf(content: bytes) -> list[tuple[str, str]]:
    try:
        import pdfplumber
    except ImportError:
        return []

    pairs: list[tuple[str, str]] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row and len(row) >= 2 and row[0] and row[1]:
                        src = _strip_prefix(str(row[0]))
                        tgt = _strip_prefix(str(row[1]))
                        if src and tgt and src.lower() not in ("word", "woord", "dutch", "english", "translation"):
                            pairs.append((src, tgt))

            if not pairs:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    pair = _split_pair(line)
                    if pair:
                        pairs.append(pair)

    return pairs


def _parse_docx(content: bytes) -> list[tuple[str, str]]:
    try:
        from docx import Document
    except ImportError:
        return []

    pairs: list[tuple[str, str]] = []
    doc = Document(io.BytesIO(content))

    # Strategy 1: tables
    for table in doc.tables:
        for row in table.rows:
            cells = [_strip_prefix(c.text) for c in row.cells]
            # Deduplicate merged cells (docx repeats cell text for merged cells)
            unique = []
            for c in cells:
                if not unique or c != unique[-1]:
                    unique.append(c)
            if len(unique) >= 2 and unique[0] and unique[1]:
                src, tgt = unique[0], unique[1]
                if src.lower() not in ("word", "woord", "dutch", "english", "translation", "term"):
                    pairs.append((src, tgt))

    if pairs:
        return pairs

    # Strategy 2: paragraphs with delimiter detection
    for para in doc.paragraphs:
        pair = _split_pair(para.text)
        if pair:
            pairs.append(pair)

    return pairs
